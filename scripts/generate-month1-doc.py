"""
Generate a client-facing Word doc for Month 1 of SEO blog publishing.
Reads content/post-queue.json, produces a polished presentation document.
"""

import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parent.parent
QUEUE_PATH = REPO / "content" / "post-queue.json"
OUT_PATH = REPO / "Gadget-Construction-SEO-Month-1.docx"

# Brand palette
BRAND_RED = RGBColor(0xCC, 0x00, 0x00)
PRIMARY = RGBColor(0x22, 0x22, 0x22)
SECONDARY = RGBColor(0x44, 0x44, 0x44)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int, color=PRIMARY) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)


def add_para(doc: Document, text: str, *, bold: bool = False, color=SECONDARY,
             size: int = 11, space_after: int = 8, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10.5)
    label_run.font.color.rgb = PRIMARY
    value_run = p.add_run(value)
    value_run.font.size = Pt(10.5)
    value_run.font.color.rgb = SECONDARY


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = SECONDARY


def format_date(iso_date: str) -> str:
    dt = datetime.fromisoformat(iso_date)
    return dt.strftime("%A, %B %-d, %Y")


def pretty_geo(geo: str) -> str:
    return {
        "SF-anchored": "San Francisco",
        "Bay Area-wide": "Bay Area (31 cities)",
        "Hyper-local": "Hyper-local (coastal)",
        "Hyper-local (coastal)": "Hyper-local (coastal)",
    }.get(geo, geo)


def build_doc() -> None:
    with open(QUEUE_PATH) as f:
        queue = json.load(f)

    month1 = queue[:4]
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ──────────── Cover ────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Q2 2026 SEO Content Plan")
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = PRIMARY
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Month 1 — Organic Traffic & Lead Generation Strategy")
    sub_run.font.size = Pt(15)
    sub_run.font.color.rgb = BRAND_RED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    meta_run = meta.add_run(
        f"Publishing Schedule: {format_date(month1[0]['scheduledDate'])} "
        f"– {format_date(month1[-1]['scheduledDate'])}"
    )
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = MUTED

    company_line = doc.add_paragraph()
    company_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_line.paragraph_format.space_after = Pt(20)
    comp_run = company_line.add_run(
        "Prepared for Gadget Construction Inc. | CA Lic #1132983"
    )
    comp_run.font.size = Pt(10)
    comp_run.font.color.rgb = MUTED
    comp_run.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ──────────── Executive Summary ────────────
    add_heading(doc, "Executive Summary", level=1)

    add_para(
        doc,
        "Over the next four weeks, Gadget Construction will publish four in-depth, "
        "search-optimized blog posts targeting commercial-intent keywords in the San "
        "Francisco Bay Area market. Each post is written for homeowners actively "
        "researching a specific construction decision — meaning they are already "
        "closer to buying than top-funnel browsers.",
        space_after=8,
    )

    total_words = sum(p["targetWordCount"] for p in month1)
    add_para(
        doc,
        f"Total output: {total_words:,} words across 4 long-form posts, averaging "
        f"{total_words // len(month1):,} words each. Every post links to the "
        "matching service page on gadgetconstructionsf.com to convert research-stage "
        "traffic into estimate requests.",
        space_after=8,
    )

    add_heading(doc, "Strategic Focus", level=2)
    summary_bullets = [
        ("Composite Decks (Posts 1, 2, 3)", "Highest-value service category — $15K–$60K+ per job. "
            "Three posts in four weeks builds topical authority on composite decking."),
        ("Foundation Underpinning (Post 4)", "High-CPC commercial keyword ($2.95 CPC low, $55 high) "
            "with urgent-intent searchers. Low competition from other contractors."),
        ("SEO Cluster Strategy", "Posts 1, 2, and 3 cross-link to each other, creating a content "
            "hub that signals expertise on composite decking to Google."),
        ("Local Signal", "Every post references specific San Francisco neighborhoods, building "
            "codes (SF DBI), and soil conditions — E-E-A-T 'Experience' signals rewarded by "
            "Google's 2026 Helpful Content update."),
    ]
    for label, description in summary_bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(6)
        label_run = p.add_run(f"{label}. ")
        label_run.bold = True
        label_run.font.size = Pt(11)
        label_run.font.color.rgb = PRIMARY
        desc_run = p.add_run(description)
        desc_run.font.size = Pt(11)
        desc_run.font.color.rgb = SECONDARY

    add_heading(doc, "Month 1 at a Glance", level=2)

    # Summary table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["#", "Publish Date", "Title", "Primary Keyword", "Words"]):
        hdr[idx].text = ""
        run = hdr[idx].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[idx], "222222")

    for i, post in enumerate(month1, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = datetime.fromisoformat(post["scheduledDate"]).strftime("%b %-d")
        row[2].text = post["title"].split(":")[0]  # short title
        row[3].text = post["primaryKeyword"]
        row[4].text = f"{post['targetWordCount']:,}"
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Page break before first post
    doc.add_page_break()

    # ──────────── Each Post ────────────
    for idx, post in enumerate(month1, 1):
        # Post header
        wk = doc.add_paragraph()
        wk.paragraph_format.space_after = Pt(4)
        wk_run = wk.add_run(f"WEEK {idx}")
        wk_run.bold = True
        wk_run.font.size = Pt(10)
        wk_run.font.color.rgb = BRAND_RED

        add_heading(doc, post["title"], level=1)

        # Meta card (table)
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = "Light Shading Accent 1"
        meta_rows = [
            ("Publishing Date", format_date(post["scheduledDate"])),
            ("Primary Keyword", post["primaryKeyword"]),
            ("Target Service", f"/services/{post['relatedService']}"),
            ("Geographic Focus", pretty_geo(post["geoFocus"])),
            ("Article Length", f"{post['targetWordCount']:,} words (~{post['targetWordCount']//250}-minute read)"),
        ]
        for row_idx, (label, value) in enumerate(meta_rows):
            row = meta_table.rows[row_idx].cells
            row[0].text = ""
            lbl_run = row[0].paragraphs[0].add_run(label)
            lbl_run.bold = True
            lbl_run.font.size = Pt(10)
            lbl_run.font.color.rgb = PRIMARY
            row[1].text = ""
            val_run = row[1].paragraphs[0].add_run(value)
            val_run.font.size = Pt(10)
            val_run.font.color.rgb = SECONDARY

        # Column widths
        for row in meta_table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(4.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Summary
        add_heading(doc, "What This Post Does", level=2)
        add_para(doc, post["excerpt"], italic=True, size=11.5)

        # Why This Post (custom analysis per post)
        add_heading(doc, "Why This Post Wins", level=2)
        rationale = POST_RATIONALE.get(post["slug"], "")
        add_para(doc, rationale, space_after=10)

        # Audience
        add_heading(doc, "Target Audience", level=2)
        audience = POST_AUDIENCE.get(post["slug"], "")
        add_para(doc, audience, space_after=10)

        # Outline
        add_heading(doc, "Content Outline", level=2)
        for section in post["outline"]:
            h2 = section["h2"]
            h2_para = doc.add_paragraph()
            h2_para.paragraph_format.space_after = Pt(3)
            h2_run = h2_para.add_run(f"▸ {h2}")
            h2_run.bold = True
            h2_run.font.size = Pt(11)
            h2_run.font.color.rgb = PRIMARY

            if "h3" in section and section["h3"]:
                for h3 in section["h3"]:
                    h3_para = doc.add_paragraph()
                    h3_para.paragraph_format.space_after = Pt(2)
                    h3_para.paragraph_format.left_indent = Inches(0.35)
                    h3_run = h3_para.add_run(f"• {h3}")
                    h3_run.font.size = Pt(10)
                    h3_run.font.color.rgb = MUTED

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # SEO focus
        add_heading(doc, "SEO Keywords Targeted", level=2)
        keyword_list = [post["primaryKeyword"]] + post["secondaryKeywords"]
        for kw in keyword_list:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(kw)
            run.font.size = Pt(10.5)
            run.font.color.rgb = SECONDARY
            if kw == post["primaryKeyword"]:
                run.bold = True

        # Internal linking
        add_heading(doc, "Site Cross-Linking Strategy", level=2)
        for link in post["internalLinks"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            anchor_run = p.add_run(f'"{link["anchor"]}" ')
            anchor_run.italic = True
            anchor_run.font.size = Pt(10.5)
            anchor_run.font.color.rgb = BRAND_RED
            url_run = p.add_run(f"→ {link['url']}")
            url_run.font.size = Pt(10)
            url_run.font.color.rgb = MUTED

        # Call to Action
        add_heading(doc, "Lead Generation CTA", level=2)
        add_para(doc, post["cta"], italic=True, size=10.5)

        if idx < len(month1):
            doc.add_page_break()

    # ──────────── Closing ────────────
    doc.add_page_break()
    add_heading(doc, "Beyond Month 1", level=1)

    add_para(
        doc,
        "Months 2 and 3 continue the strategy with six additional long-form posts "
        "covering stucco repair, dry rot diagnosis, siding replacement, hillside "
        "deck construction, and a coastal Bay Area exterior-repairs authority piece. "
        "The full 10-post plan forms three interlocking topic clusters — composite "
        "decking, foundation/underpinning, and exterior repairs — each pointing "
        "readers toward their matching service page on gadgetconstructionsf.com.",
        space_after=10,
    )

    add_heading(doc, "Measurement", level=2)
    measurement = [
        "Organic search traffic to the domain (Google Search Console)",
        "Keyword ranking positions for each target keyword (tracked weekly)",
        "Form submissions attributed to organic traffic (Google Analytics 4)",
        "Phone call leads attributed to organic sessions (CallRail dynamic number insertion)",
        "Closed jobs originating from each blog post (manual attribution at estimate stage)",
    ]
    for item in measurement:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = SECONDARY

    add_heading(doc, "Publishing Workflow", level=2)
    workflow = [
        "Each week's post is drafted automatically from a pre-approved content brief",
        "Raul reviews the draft on Friday or over the weekend before publishing",
        "Final post publishes to gadgetconstructionsf.com/blog on its scheduled Monday",
        "Google is notified via sitemap and indexing request",
        "Performance is reviewed monthly with adjustments to the following month's plan",
    ]
    for item in workflow:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Footer line
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Gadget Construction Inc. | gadgetconstructionsf.com | CA Lic #1132983"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    footer_run.italic = True

    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")


# Per-post custom rationale for "Why This Post Wins"
POST_RATIONALE = {
    "composite-deck-cost-san-francisco": (
        "Cost is the single most-searched question before a homeowner commits to a "
        "deck project. Searches like 'composite deck cost san francisco' have high "
        "commercial intent — these homeowners are actively budgeting, not casually "
        "browsing. Most existing ranking content is either national-average data or "
        "contractor directory listings. By publishing Gadget's actual numbers from "
        "500+ Bay Area projects, this post gives homeowners what they came for and "
        "establishes Gadget as the authoritative local source."
    ),
    "trex-vs-timbertech-vs-fiberon-bay-area-2026": (
        "Brand-comparison searches are decision-stage queries. Homeowners asking "
        "'Trex vs TimberTech vs Fiberon' have already decided they want composite "
        "decking and are choosing between brands. Most ranking content is from the "
        "brand manufacturers themselves (biased) or retailers (shallow). An "
        "independent, contractor-written comparison — backed by installations in "
        "all three brands across Bay Area weather conditions — is rare and highly "
        "ranked by Google."
    ),
    "composite-vs-wood-decking-bay-area-2026": (
        "This post captures a wider audience — homeowners who haven't yet chosen "
        "composite. It educates them on the total cost of ownership math that "
        "favors composite in the Bay Area's fog-belt climate, then converts them "
        "to the composite service page. The post also honestly acknowledges when "
        "wood still wins (for trust) — a counterintuitive framing Google rewards "
        "under the Helpful Content guidelines."
    ),
    "foundation-underpinning-cost-san-francisco": (
        "Foundation underpinning is an urgent-need service. Homeowners searching "
        "'foundation underpinning cost' already have a problem they know needs "
        "fixing — the question is only who and how much. This keyword has low "
        "competition (most contractors don't produce educational content on method "
        "comparisons) and very high commercial intent. The post's decision-framework "
        "format — helical vs push vs slabjacking — makes it uniquely useful and "
        "differentiates Gadget as an expert rather than a quote-and-close contractor."
    ),
}

# Per-post target audience
POST_AUDIENCE = {
    "composite-deck-cost-san-francisco": (
        "San Francisco homeowners actively planning a deck renovation, typically "
        "in neighborhoods with outdoor space (Noe Valley, Bernal Heights, Twin "
        "Peaks, Glen Park, Diamond Heights, Forest Hill). Primarily mid-to-high "
        "income households ($200K+ HHI) evaluating a $15,000–$60,000 project. "
        "Stage: budgeting before getting estimates."
    ),
    "trex-vs-timbertech-vs-fiberon-bay-area-2026": (
        "Bay Area homeowners who have already decided on composite decking and "
        "are choosing between brands. Typically contractor-ready — within 30 days "
        "of signing a contract. Often researching specific warranty claims, "
        "color fade resistance, and Bay Area coastal durability."
    ),
    "composite-vs-wood-decking-bay-area-2026": (
        "Bay Area homeowners 3–6 months from starting a deck project, still "
        "evaluating materials. Often influenced by an architect or designer "
        "recommendation. Wider funnel than cost posts — captures research-stage "
        "traffic and educates them toward composite, which is Gadget's "
        "higher-margin service."
    ),
    "foundation-underpinning-cost-san-francisco": (
        "San Francisco homeowners with active foundation problems — visible "
        "cracks, sloping floors, sticking doors. Often in Marina District (fill "
        "soil), Sunset (clay soil), or hillside neighborhoods (Bernal Heights, "
        "Twin Peaks). High-urgency buyers evaluating a $15,000–$50,000+ repair. "
        "Frequently have an engineer's report in hand."
    ),
}


if __name__ == "__main__":
    build_doc()
