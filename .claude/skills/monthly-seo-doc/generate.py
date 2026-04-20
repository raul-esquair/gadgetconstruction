"""
Generate a client-facing Word doc for the next N blog posts.

Parameterized version of the month-1 script. Reads a payload JSON that lists
which posts to include and provides custom client rationale + target audience
for each. Produces a polished .docx in the repo root.

Usage:
  python3 generate.py --payload /tmp/skill-payload.json --output ./Month-2.docx

Payload JSON structure:
{
  "monthNumber": 2,
  "dateRangeLabel": "May 4 – May 25, 2026",           # optional
  "posts": [
    {
      "brief": { ...full post brief from post-queue.json... },
      "clientRationale": "3-4 sentence SEO rationale for this post",
      "targetAudience":  "3-4 sentence audience description"
    },
    ...
  ]
}
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Brand palette (from CLAUDE.md design system)
BRAND_RED = RGBColor(0xCC, 0x00, 0x00)
PRIMARY = RGBColor(0x22, 0x22, 0x22)
SECONDARY = RGBColor(0x44, 0x44, 0x44)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int, color=PRIMARY) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(22)
    elif level == 2:
        run.font.size = Pt(16)
    elif level == 3:
        run.font.size = Pt(13)


def add_para(doc: Document, text: str, *, bold: bool = False, color=SECONDARY,
             size: float = 11, space_after: int = 8, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def format_date(iso_date: str) -> str:
    dt = datetime.fromisoformat(iso_date)
    return dt.strftime("%A, %B %-d, %Y")


def pretty_geo(geo: str) -> str:
    return {
        "SF-anchored": "San Francisco",
        "Bay Area-wide": "Bay Area (31 cities)",
        "Hyper-local": "Hyper-local (coastal)",
        "Hyper-local (coastal)": "Hyper-local (coastal)",
    }.get(geo, geo)


def ordinal(n: int) -> str:
    if 10 <= n % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suffix}"


def build_doc(payload: dict, out_path: Path) -> None:
    month_num = payload["monthNumber"]
    posts = payload["posts"]
    if not posts:
        raise ValueError("Payload contains no posts")

    briefs = [p["brief"] for p in posts]
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ──────────── Cover ────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Q2 2026 SEO Content Plan")
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = PRIMARY
    title_run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    month_label = (
        f"Month {month_num} — Organic Traffic & Lead Generation Strategy"
    )
    sub_run = subtitle.add_run(month_label)
    sub_run.font.size = Pt(15)
    sub_run.font.color.rgb = BRAND_RED

    date_label = payload.get(
        "dateRangeLabel",
        f"{format_date(briefs[0]['scheduledDate'])} – {format_date(briefs[-1]['scheduledDate'])}",
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    meta_run = meta.add_run(f"Publishing Schedule: {date_label}")
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = MUTED

    company_line = doc.add_paragraph()
    company_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_line.paragraph_format.space_after = Pt(20)
    comp_run = company_line.add_run(
        "Prepared for Gadget Construction Inc. | CA Lic #1132983"
    )
    comp_run.font.size = Pt(10)
    comp_run.font.color.rgb = MUTED
    comp_run.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ──────────── Executive Summary ────────────
    add_heading(doc, "Executive Summary", level=1)

    total_words = sum(b["targetWordCount"] for b in briefs)
    add_para(
        doc,
        f"This is the {ordinal(month_num)} month of Gadget Construction's SEO blog "
        f"publishing program. Over the next {len(briefs)} weeks, we will publish "
        f"{len(briefs)} in-depth, search-optimized blog posts targeting "
        "commercial-intent keywords in the San Francisco Bay Area market. Each "
        "post is written for homeowners actively researching a specific "
        "construction decision — meaning they are already closer to buying than "
        "top-funnel browsers.",
        space_after=8,
    )

    add_para(
        doc,
        f"Total output: {total_words:,} words across {len(briefs)} long-form "
        f"posts, averaging {total_words // len(briefs):,} words each. Every post "
        "links to the matching service page on gadgetconstructionsf.com to "
        "convert research-stage traffic into estimate requests.",
        space_after=8,
    )

    # Executive summary bullets pulled from payload if provided, otherwise skip
    exec_bullets = payload.get("executiveSummaryBullets")
    if exec_bullets:
        add_heading(doc, "Strategic Focus", level=2)
        for label, description in exec_bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(6)
            label_run = p.add_run(f"{label}. ")
            label_run.bold = True
            label_run.font.size = Pt(11)
            label_run.font.color.rgb = PRIMARY
            desc_run = p.add_run(description)
            desc_run.font.size = Pt(11)
            desc_run.font.color.rgb = SECONDARY

    add_heading(doc, f"Month {month_num} at a Glance", level=2)

    # Summary table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for idx, text in enumerate(["#", "Publish Date", "Title", "Primary Keyword", "Words"]):
        hdr[idx].text = ""
        run = hdr[idx].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[idx], "222222")

    start_index_in_month = 1
    for i, brief in enumerate(briefs, start=start_index_in_month):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = datetime.fromisoformat(brief["scheduledDate"]).strftime("%b %-d")
        row[2].text = brief["title"].split(":")[0]
        row[3].text = brief["primaryKeyword"]
        row[4].text = f"{brief['targetWordCount']:,}"
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Page break before first post
    doc.add_page_break()

    # ──────────── Each Post ────────────
    for idx, post in enumerate(posts, 1):
        brief = post["brief"]
        client_rationale = post.get("clientRationale", "")
        target_audience = post.get("targetAudience", "")

        # Post header
        wk = doc.add_paragraph()
        wk.paragraph_format.space_after = Pt(4)
        wk_run = wk.add_run(f"WEEK {idx}")
        wk_run.bold = True
        wk_run.font.size = Pt(10)
        wk_run.font.color.rgb = BRAND_RED

        add_heading(doc, brief["title"], level=1)

        # Meta card
        meta_table = doc.add_table(rows=5, cols=2)
        meta_table.style = "Light Shading Accent 1"
        meta_rows = [
            ("Publishing Date", format_date(brief["scheduledDate"])),
            ("Primary Keyword", brief["primaryKeyword"]),
            ("Target Service", f"/services/{brief['relatedService']}"),
            ("Geographic Focus", pretty_geo(brief["geoFocus"])),
            (
                "Article Length",
                f"{brief['targetWordCount']:,} words "
                f"(~{brief['targetWordCount'] // 250}-minute read)",
            ),
        ]
        for row_idx, (label, value) in enumerate(meta_rows):
            row = meta_table.rows[row_idx].cells
            row[0].text = ""
            lbl_run = row[0].paragraphs[0].add_run(label)
            lbl_run.bold = True
            lbl_run.font.size = Pt(10)
            lbl_run.font.color.rgb = PRIMARY
            row[1].text = ""
            val_run = row[1].paragraphs[0].add_run(value)
            val_run.font.size = Pt(10)
            val_run.font.color.rgb = SECONDARY

        for row in meta_table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(4.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Summary
        add_heading(doc, "What This Post Does", level=2)
        add_para(doc, brief["excerpt"], italic=True, size=11.5)

        # Why This Post Wins
        if client_rationale:
            add_heading(doc, "Why This Post Wins", level=2)
            add_para(doc, client_rationale, space_after=10)

        # Audience
        if target_audience:
            add_heading(doc, "Target Audience", level=2)
            add_para(doc, target_audience, space_after=10)

        # Outline
        add_heading(doc, "Content Outline", level=2)
        for section in brief["outline"]:
            h2 = section["h2"]
            h2_para = doc.add_paragraph()
            h2_para.paragraph_format.space_after = Pt(3)
            h2_run = h2_para.add_run(f"▸ {h2}")
            h2_run.bold = True
            h2_run.font.size = Pt(11)
            h2_run.font.color.rgb = PRIMARY

            if "h3" in section and section["h3"]:
                for h3 in section["h3"]:
                    h3_para = doc.add_paragraph()
                    h3_para.paragraph_format.space_after = Pt(2)
                    h3_para.paragraph_format.left_indent = Inches(0.35)
                    h3_run = h3_para.add_run(f"• {h3}")
                    h3_run.font.size = Pt(10)
                    h3_run.font.color.rgb = MUTED

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # SEO focus
        add_heading(doc, "SEO Keywords Targeted", level=2)
        keyword_list = [brief["primaryKeyword"]] + brief.get("secondaryKeywords", [])
        for kw in keyword_list:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(kw)
            run.font.size = Pt(10.5)
            run.font.color.rgb = SECONDARY
            if kw == brief["primaryKeyword"]:
                run.bold = True

        # Internal linking
        add_heading(doc, "Site Cross-Linking Strategy", level=2)
        for link in brief.get("internalLinks", []):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            anchor_run = p.add_run(f'"{link["anchor"]}" ')
            anchor_run.italic = True
            anchor_run.font.size = Pt(10.5)
            anchor_run.font.color.rgb = BRAND_RED
            url_run = p.add_run(f"→ {link['url']}")
            url_run.font.size = Pt(10)
            url_run.font.color.rgb = MUTED

        # CTA
        add_heading(doc, "Lead Generation CTA", level=2)
        add_para(doc, brief["cta"], italic=True, size=10.5)

        if idx < len(posts):
            doc.add_page_break()

    # ──────────── Closing ────────────
    doc.add_page_break()
    add_heading(doc, f"Beyond Month {month_num}", level=1)

    closing_text = payload.get(
        "closingParagraph",
        (
            "The publishing program continues with additional long-form posts "
            "covering adjacent topics in the same service categories. Each post "
            "forms part of an interlocking topic cluster, reinforcing Gadget's "
            "expertise across composite decking, foundations, and exterior "
            "repairs. Every post points readers toward its matching service page "
            "on gadgetconstructionsf.com."
        ),
    )
    add_para(doc, closing_text, space_after=10)

    add_heading(doc, "Measurement", level=2)
    measurement = [
        "Organic search traffic to the domain (Google Search Console)",
        "Keyword ranking positions for each target keyword (tracked weekly)",
        "Form submissions attributed to organic traffic (Google Analytics 4)",
        "Phone call leads attributed to organic sessions (CallRail dynamic number insertion)",
        "Closed jobs originating from each blog post (manual attribution at estimate stage)",
    ]
    for item in measurement:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = SECONDARY

    add_heading(doc, "Publishing Workflow", level=2)
    workflow = [
        "Each week's post is drafted automatically from a pre-approved content brief",
        "Raul reviews the draft on Friday or over the weekend before publishing",
        "Final post publishes to gadgetconstructionsf.com/blog on its scheduled Monday",
        "Google is notified via sitemap and indexing request",
        "Performance is reviewed monthly with adjustments to the following month's plan",
    ]
    for item in workflow:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = SECONDARY

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        "Gadget Construction Inc. | gadgetconstructionsf.com | CA Lic #1132983"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    footer_run.italic = True

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--payload", required=True, type=Path,
                        help="Path to JSON payload with posts + month metadata")
    parser.add_argument("--output", required=True, type=Path,
                        help="Output .docx path")
    args = parser.parse_args()

    if not args.payload.exists():
        print(f"ERROR: payload file not found: {args.payload}", file=sys.stderr)
        sys.exit(1)

    with open(args.payload) as f:
        payload = json.load(f)

    build_doc(payload, args.output)
    print(f"Generated: {args.output}")


if __name__ == "__main__":
    main()
